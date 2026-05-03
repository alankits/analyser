"""Text extraction from PDF and DOCX files."""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_pdf(path: Path) -> str:
    """Extract text from a PDF using pdfplumber. Falls back to OCR if text is sparse."""
    import pdfplumber  # type: ignore

    pages_text: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages_text.append(text)
                logger.debug("PDF page %d extracted %d chars", i + 1, len(text))
    except Exception as exc:
        logger.warning("pdfplumber extraction failed: %s", exc)
        return ""

    full_text = "\n".join(pages_text)

    # If extracted text is too sparse, it's likely a scanned PDF — trigger OCR
    if len(full_text.strip()) < 200:
        logger.info("Sparse text detected (%d chars), triggering OCR", len(full_text))
        from modules.ocr_extractor import extract_text_via_ocr

        ocr_text = extract_text_via_ocr(path)
        return ocr_text if ocr_text.strip() else full_text

    return full_text


def extract_text_from_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document  # type: ignore

    try:
        doc = Document(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        all_text = "\n".join(paragraphs + table_texts)
        logger.info("DOCX extracted %d chars", len(all_text))
        return all_text
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""


def extract_text(path: Path, file_type: str) -> str:
    """
    Dispatch to the correct extractor based on file_type ('pdf' or 'docx').

    Returns:
        Extracted plain text string (may be empty if extraction fails).
    """
    if file_type == "pdf":
        return extract_text_from_pdf(path)
    elif file_type == "docx":
        return extract_text_from_docx(path)
    else:
        logger.error("Unknown file type for extraction: %s", file_type)
        return ""
